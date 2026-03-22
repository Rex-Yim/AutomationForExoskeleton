#!/usr/bin/env python3
"""Build docs/1155176974_Yim_Fu_Chong_MAEG4998_Final_Report.docx from docs/Final_Report.txt."""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

HEADING1 = frozenset(
    {
        "Acknowledgments",
        "Abstract",
        "Problem Description",
        "Background Knowledge and Literature Review",
        "Proposed Solution Methods",
        "Results and Discussions",
        "Conclusions",
        "References",
    }
)


def clear_document_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        body.remove(child)


def normalize_line(s: str) -> str:
    s = s.replace("vs.\\ ", "vs. ")
    s = s.replace("``", '"').replace("''", '"')
    s = s.replace('Tr\\"oster', "Tröster")
    return s.strip()


def _next_non_empty(lines: list[str], start: int) -> int:
    j = start
    while j < len(lines) and not lines[j].strip():
        j += 1
    return j


def is_heading2(lines: list[str], i: int) -> bool:
    line = lines[i].strip()
    if not line or line in HEADING1:
        return False
    if line.startswith(("-", "[", "%")):
        return False
    if i + 1 < len(lines) and lines[i + 1].strip():
        return False
    j = _next_non_empty(lines, i + 1)
    if j >= len(lines):
        return False
    nxt = lines[j].strip()
    if len(nxt) > 100:
        return True
    k = _next_non_empty(lines, j + 1)
    if k >= len(lines):
        return False
    return len(lines[k].strip()) > 100


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    txt_path = root / "docs" / "Final_Report.txt"
    desktop = Path.home() / "Desktop" / "1155176974_Yim_Fu_Chong_MAEG4998_Final_Report.docx"
    out_path = root / "docs" / "1155176974_Yim_Fu_Chong_MAEG4998_Final_Report.docx"

    if not txt_path.is_file():
        print(f"Missing {txt_path}", file=sys.stderr)
        return 1

    if desktop.is_file():
        shutil.copy2(desktop, out_path)
        doc = Document(str(out_path))
        clear_document_body(doc)
    else:
        doc = Document()
        out_path.parent.mkdir(parents=True, exist_ok=True)

    raw = txt_path.read_text(encoding="utf-8")
    lines = raw.splitlines()

    i = 0
    while i < len(lines) and "=" not in lines[i]:
        i += 1
    i += 1

    while i < len(lines) and lines[i].strip() != "Acknowledgments":
        t = normalize_line(lines[i])
        if t:
            p = doc.add_paragraph(t)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1

    last_was_references = False
    while i < len(lines):
        line = lines[i]
        t = normalize_line(line)

        if not t:
            i += 1
            continue

        if t == "References" and last_was_references:
            i += 1
            continue

        if t in HEADING1:
            doc.add_heading(t, level=1)
            last_was_references = t == "References"
            i += 1
            continue

        last_was_references = False

        if t.startswith("[Figure file:"):
            m = re.search(r"\[Figure file:\s*([^\]]+)\]", t)
            if m:
                fname = m.group(1).strip()
                img = root / "results" / fname
                if img.is_file():
                    try:
                        doc.add_picture(str(img), width=Inches(5.8))
                    except OSError:
                        doc.add_paragraph(t)
                else:
                    doc.add_paragraph(t)
            i += 1
            continue

        if t.startswith("- "):
            try:
                p = doc.add_paragraph(t[2:].strip(), style="List Bullet")
            except KeyError:
                p = doc.add_paragraph(f"• {t[2:].strip()}")
            i += 1
            continue

        if is_heading2(lines, i):
            doc.add_heading(t, level=2)
            i += 1
            continue

        doc.add_paragraph(t)
        i += 1

    doc.save(str(out_path))
    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
